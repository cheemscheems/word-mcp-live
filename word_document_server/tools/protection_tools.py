"""
Protection tools for Word Document Server.

These tools handle document protection features such as
password protection, OOXML restricted editing, and signature blocks.
"""
import datetime
import io
import json
import os
import sys
from typing import List, Optional

from docx import Document
from docx.oxml.shared import OxmlElement, qn
import msoffcrypto

from word_document_server.utils.file_utils import (
    check_file_writeable,
    ensure_docx_extension,
    get_file_lock,
)
from word_document_server.core.document_protection import (
    add_restricted_editing_protection,
    remove_restricted_editing_protection,
    has_document_protection,
)


async def protect_document(
    filename: str,
    password: str,
    confirm_password: str = None,
) -> str:
    """Add password protection (encryption) to a Word document.

    Uses msoffcrypto-tool to apply real OOXML-level encryption.  The
    password is **not** stored anywhere -- it is used only for the
    encryption operation and then discarded.  If the password is lost
    the document cannot be recovered.

    Args:
        filename:         Path to the Word document.
        password:         Password to protect the document with.
        confirm_password: Optional -- must match *password* when provided.
                          Use this to let the user confirm the password
                          before encryption.
    """
    filename = ensure_docx_extension(filename)

    if not os.path.exists(filename):
        return json.dumps({"success": False, "error": f"Document {filename} does not exist"})

    if confirm_password is not None and password != confirm_password:
        return json.dumps({
            "success": False,
            "error": "两次输入的密码不一致，请重新输入。",
        })

    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return json.dumps({"success": False, "error": f"Cannot protect document: {error_message}"})

    try:
        async with get_file_lock(filename):
            with open(filename, "rb") as infile:
                original_data = infile.read()

            file = msoffcrypto.OfficeFile(io.BytesIO(original_data))
            file.load_key(password=password)

            encrypted_data_io = io.BytesIO()
            file.encrypt(password=password, outfile=encrypted_data_io)

            with open(filename, "wb") as outfile:
                outfile.write(encrypted_data_io.getvalue())

        return json.dumps({
            "success": True,
            "document": os.path.basename(filename),
            "encrypted": True,
            "password_warning": (
                "⚠️ 密码不会被保存。请务必牢记此密码，否则将无法解密文档。"
                "建议将密码记录在安全的地方。"
            ),
            "user_guidance": (
                "文档已加密。密码仅在此次会话中有效，不会被存储。"
                "请牢记密码，丢失后无法找回。"
            ),
        })

    except Exception as e:
        try:
            if "original_data" in locals():
                with open(filename, "wb") as outfile:
                    outfile.write(original_data)
                return json.dumps({
                    "success": False,
                    "error": f"Failed to encrypt: {str(e)}. Original file restored.",
                })
        except Exception as restore_e:
            pass
        return json.dumps({"success": False, "error": str(e)})


async def add_restricted_editing(
    filename: str,
    password: str,
    editable_sections: List[str],
) -> str:
    """Add restricted editing to a Word document.

    Inserts a ``<w:documentProtection>`` element into ``word/settings.xml``
    so that Word's own Restrict Editing pane enforces the restriction.

    The *editable_sections* parameter is mapped to one of Word's four
    protection modes:

    +------------------------------------+--------------+
    | ``editable_sections``              | Mode         |
    +------------------------------------+--------------+
    | ``[]`` or ``["none"]``             | readOnly     |
    | ``["comments"]``                   | comments     |
    | ``["trackedChanges"]`` / revisions | trackedChanges |
    | ``["forms"]`` / ``["fillIn"]``     | forms        |
    | anything else                      | readOnly     |
    +------------------------------------+--------------+

    Args:
        filename:          Path to the Word document.
        password:          Password required to change or remove protection.
        editable_sections: Mapped to protection mode (see table above).
    """
    filename = ensure_docx_extension(filename)

    if not os.path.exists(filename):
        return json.dumps({"success": False, "error": f"Document {filename} does not exist"})

    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return json.dumps({"success": False, "error": f"Cannot protect document: {error_message}"})

    # Map editable_sections → OOXML protection mode
    _mode_map = {
        "none": "readOnly",
        "comments": "comments",
        "trackedchanges": "trackedChanges",
        "revisions": "trackedChanges",
        "forms": "forms",
        "fillin": "forms",
    }
    if editable_sections:
        key = editable_sections[0].strip().lower()
        edit_mode = _mode_map.get(key, "readOnly")
    else:
        edit_mode = "readOnly"

    try:
        async with get_file_lock(filename):
            doc = Document(filename)
            add_restricted_editing_protection(doc, edit_mode, password)
            doc.save(filename)

        return json.dumps({
            "success": True,
            "document": os.path.basename(filename),
            "edit_mode": edit_mode,
            "message": f"Document protected with {edit_mode} editing.",
        })

    except Exception as e:
        return json.dumps({"success": False, "error": str(e)})


async def add_signature_block(
    filename: str,
    signer_name: str,
    title: str = None,
    reason: str = None,
    location: str = None,
    show_date: bool = True,
) -> str:
    """Add a visible signature block to a Word document.

    - **Cross-platform (python-docx):** inserts a formal text-based
      signature block with a separator line, signer name, title, date,
      location, and a signature placeholder line.
    - **Windows + COM (when Word is running):** additionally creates a
      native Word Signature Line shape (the ``doc.Signatures`` COM
      object) that the recipient can double-click to sign.

    .. note::
       The COM signature line is a **placeholder** only — it does not
       perform cryptographic signing.  To complete the signature, the
       recipient opens the document in Word and double-clicks the
       signature line, then either types a signature or uses a
       certificate.

    Args:
        filename:     Path to the Word document.
        signer_name:  Name of the person signing.
        title:        Optional job title of the signer.
        reason:       Optional reason for signing.
        location:     Optional signing location (city / office).
        show_date:    Whether to include the current date (default True).
    """
    filename = ensure_docx_extension(filename)

    if not os.path.exists(filename):
        return json.dumps({"success": False, "error": f"Document {filename} does not exist"})

    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return json.dumps({"success": False, "error": f"Cannot add signature: {error_message}"})

    has_com_line = False

    try:
        async with get_file_lock(filename):
            doc = Document(filename)

            # ── Cross-platform text signature block ──────────────────────
            # Separator line (bottom border on an empty paragraph)
            sep = doc.add_paragraph()
            pPr = sep._element.get_or_add_pPr()
            from docx.oxml.shared import OxmlElement
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "12")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "000000")
            pBdr.append(bottom)
            pPr.append(pBdr)

            doc.add_paragraph("")  # spacing

            sig_para = doc.add_paragraph()
            sig_para.add_run(f"{signer_name}").bold = True
            if title:
                sig_para.add_run(f"\n{title}")

            detail_line = ""
            if location:
                detail_line += f"{location}"
            if show_date:
                date_str = datetime.datetime.now().strftime("%Y-%m-%d")
                if detail_line:
                    detail_line += f"，{date_str}"
                else:
                    detail_line = date_str
            if detail_line:
                sig_para.add_run(f"\n{detail_line}")

            if reason:
                sig_para.add_run(f"\n签署理由：{reason}")

            doc.add_paragraph("")
            sig_line = doc.add_paragraph()
            sig_line.add_run("签名：_______________________________")
            doc.add_paragraph("")

            # ── Windows COM: native Signature Line shape ────────────────
            if sys.platform == "win32":
                try:
                    from word_document_server.core.word_com import (
                        get_word_app,
                        find_document,
                        undo_record,
                    )

                    app = get_word_app()
                    wdoc = find_document(app, filename)

                    with undo_record(app, "MCP: Add Signature Block"):
                        rng = wdoc.Content
                        rng.Collapse(0)  # wdCollapseEnd

                        sig = wdoc.Signatures.Add(rng)
                        sig.Signer = signer_name
                        if title:
                            sig.SignerTitle = title
                        if reason:
                            sig.Instructions = reason
                        sig.ShowSignDate = show_date

                        wdoc.Save()
                    has_com_line = True
                except Exception:
                    has_com_line = False
            else:
                has_com_line = False

            doc.save(filename)

        return json.dumps({
            "success": True,
            "document": os.path.basename(filename),
            "signer": signer_name,
            "has_com_signature_line": has_com_line,
            "has_text_block": True,
            "user_guidance": (
                "签名行已添加。请打开 Word 文档，双击签名行图标"
                " → 在弹窗中点击「签名」按钮"
                " → 输入您的姓名作为手写签名（或使用数字证书）。"
                "完成后保存文档，签名即生效。"
            ),
        })

    except Exception as e:
        return json.dumps({"success": False, "error": str(e)})


async def verify_document(filename: str, password: Optional[str] = None) -> str:
    """Verify document protection status.

    Checks for:
    - OOXML restricted editing (``<w:documentProtection>`` in settings.xml)
    - msoffcrypto-tool encryption
    - COM signature lines (Windows only)

    Args:
        filename: Path to the Word document.
        password: Optional password to try against encryption.
    """
    filename = ensure_docx_extension(filename)

    if not os.path.exists(filename):
        return json.dumps({"success": False, "error": f"Document {filename} does not exist"})

    findings = []

    try:
        doc = Document(filename)

        # 1. Check OOXML restricted editing
        is_protected, edit_mode = has_document_protection(doc)
        if is_protected:
            findings.append(f"restricted_editing: enabled (mode={edit_mode})")
        else:
            findings.append("restricted_editing: none")

        # 2. Check COM signature lines
        if sys.platform == "win32":
            try:
                from word_document_server.core.word_com import get_word_app, find_document

                app = get_word_app()
                wdoc = find_document(app, filename)
                sig_count = wdoc.Signatures.Count
                if sig_count > 0:
                    findings.append(f"com_signature_lines: {sig_count}")
                else:
                    findings.append("com_signature_lines: none")
            except Exception:
                findings.append("com_signature_lines: unavailable (Word not running?)")

    except Exception as e:
        findings.append(f"document_read_error: {str(e)}")

    # 3. Check msoffcrypto encryption
    try:
        with open(filename, "rb") as f:
            office_file = msoffcrypto.OfficeFile(f)
            if office_file.is_encrypted():
                findings.append("encryption: password-protected")
                if password:
                    try:
                        office_file.load_key(password=password)
                        # load_key alone does NOT verify the password -- only decrypt()
                        # will throw InvalidKeyError for a wrong password.
                        verify_buf = io.BytesIO()
                        office_file.decrypt(outfile=verify_buf)
                        findings.append("encryption_password: correct")
                    except msoffcrypto.exceptions.InvalidKeyError:
                        findings.append("encryption_password: incorrect")
                    except Exception:
                        findings.append("encryption_password: check_error")
            else:
                findings.append("encryption: none")
    except Exception as e:
        findings.append(f"encryption_check_error: {str(e)}")

    return json.dumps({
        "success": True,
        "document": os.path.basename(filename),
        "findings": findings,
    })


async def unprotect_document(filename: str, password: str) -> str:
    """Remove password protection (decrypt) from a Word document.

    Uses msoffcrypto-tool to decrypt the OOXML-level encryption applied
    by :func:`protect_document`.

    Args:
        filename: Path to the Word document.
        password: Password that was used to encrypt the document.
    """
    filename = ensure_docx_extension(filename)

    if not os.path.exists(filename):
        return json.dumps({"success": False, "error": f"Document {filename} does not exist"})

    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return json.dumps({"success": False, "error": f"Cannot modify document: {error_message}"})

    try:
        async with get_file_lock(filename):
            with open(filename, "rb") as infile:
                encrypted_data = infile.read()

            file = msoffcrypto.OfficeFile(io.BytesIO(encrypted_data))
            file.load_key(password=password)

            decrypted_data_io = io.BytesIO()
            file.decrypt(outfile=decrypted_data_io)

            with open(filename, "wb") as outfile:
                outfile.write(decrypted_data_io.getvalue())

        return json.dumps({
            "success": True,
            "document": os.path.basename(filename),
            "decrypted": True,
        })

    except msoffcrypto.exceptions.InvalidKeyError:
        return json.dumps({"success": False, "error": "密码不正确。"})
    except msoffcrypto.exceptions.InvalidFormatError:
        return json.dumps({
            "success": False,
            "error": "文件未加密或不是受支持的 Office 格式。",
        })
    except Exception as e:
        try:
            if "encrypted_data" in locals():
                with open(filename, "wb") as outfile:
                    outfile.write(encrypted_data)
                return json.dumps({
                    "success": False,
                    "error": f"解密失败：{str(e)}。已恢复原始文件。",
                })
        except Exception as restore_e:
            pass
        return json.dumps({"success": False, "error": str(e)})
